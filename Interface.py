import streamlit as st
import pandas as pd
import plotly.express as px  # Adicionado para os gráficos do Dashboard
from datetime import datetime
import os
import io
import unicodedata
from docx import Document
from docx.shared import Pt
# from docx2pdf import convert # Nota: convert costuma exigir Word instalado, use com cautela em servidores
from num2words import num2words

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Sistema de Gestão", layout="wide")

# --- CARREGAMENTO INICIAL ---
try:
    dados_produtos = pd.read_csv("Base de Dados.csv", sep=";", encoding="latin-1")
    dados_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";", encoding="latin-1")
except FileNotFoundError:
    st.error("Bases de dados não encontradas. Verifique os arquivos CSV.")

# --- INICIALIZAÇÃO DO ESTADO ---
if "pagina" not in st.session_state:
    st.session_state.pagina = "Dashboard" # Dashboard definido como página inicial

# --- ESTILIZAÇÃO CSS ---
st.markdown("""
<style>
    /* Botões do Menu Lateral */
    div[data-testid="stSidebar"] .stButton button {
        width: 100%;
        border-radius: 5px;
        height: 3em;
        background-color: #262730;
        border: 1px solid #464b5d;
        text-align: left;
        padding-left: 15px;
        margin-bottom: 5px;
        color: #FAFAFA;
    }
    div[data-testid="stSidebar"] .stButton button:hover {
        border-color: #ff4b4b;
        color: #ff4b4b;
    }
    /* Destaque Ativo */
    div[data-testid="stSidebar"] .active-btn button {
        background-color: #ff4b4b !important;
        color: white !important;
        border: none !important;
        font-weight: bold;
    }
    /* Estilo do Título do Expander */
    .st-emotion-cache-p4mowd {
        font-weight: bold;
        color: #ffffff;
    }
</style>
""", unsafe_allow_html=True)

# --- FUNÇÃO AUXILIAR DO MENU ---
def gerar_botao_menu(label):
    if st.session_state.pagina == label:
        st.markdown('<div class="active-btn">', unsafe_allow_html=True)
        if st.button(label, key=f"btn_{label}"):
            st.session_state.pagina = label
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        if st.button(label, key=f"btn_{label}"):
            st.session_state.pagina = label
            st.rerun()

# --- MENU LATERAL (NAVEGAÇÃO) ---
with st.sidebar:
    st.title("ERP")
    
    # Botão de Dashboard no Topo (Destaque Principal)
    gerar_botao_menu("Dashboard")
    
    st.divider()

    with st.expander("Pedidos", expanded=(st.session_state.pagina in ["Criar Pedido", "Consultar Pedido"])):
        gerar_botao_menu("Criar Pedido")
        gerar_botao_menu("Consultar Pedido")

    with st.expander("Produtos", expanded=(st.session_state.pagina in ["Cadastrar Produto", "Consultar Produto"])):
        gerar_botao_menu("Cadastrar Produto")
        gerar_botao_menu("Consultar Produto")

    with st.expander("Pessoas", expanded=(st.session_state.pagina in ["Cadastrar Pessoa", "Consultar Pessoa"])):
        gerar_botao_menu("Cadastrar Pessoa")
        gerar_botao_menu("Consultar Pessoa")

    with st.expander("Relatórios", expanded=(st.session_state.pagina == "Formalizacao")):
        gerar_botao_menu("Formalizacao")

# Define a variável para controle de fluxo
pagina = st.session_state.pagina

# --- A PARTIR DAQUI COMEÇAM OS IF PAGINA == ... ---

if pagina == "Cadastrar Produto":
    st.title("Cadastro de Produtos")

    # --- 1. LÓGICA DE CARREGAMENTO DE FORNECEDORES ---
    # Começa com uma lista padrão caso o arquivo não exista ou esteja vazio
    lista_fornecedores = ["Outros"] 
    arquivo_pessoas = "Base_Pessoas.csv"

    if os.path.exists(arquivo_pessoas):
        try:
            # Tenta ler UTF-8, se falhar tenta Latin-1
            try:
                df_pessoas = pd.read_csv(arquivo_pessoas, sep=";")
            except UnicodeDecodeError:
                df_pessoas = pd.read_csv(arquivo_pessoas, sep=";", encoding="latin1")
            
            # Remove espaços em branco dos nomes das colunas (segurança)
            df_pessoas.columns = df_pessoas.columns.str.strip()

            # Se as colunas existirem, filtra e preenche a lista
            if "nome_razao" in df_pessoas.columns and "categoria" in df_pessoas.columns:
                filtro = df_pessoas[
                    df_pessoas["categoria"].astype(str).str.contains("Fornecedor|Ambos", case=False, na=False)
                ]
                nomes_banco = sorted(filtro["nome_razao"].unique().tolist())
                
                if nomes_banco:
                    # Adiciona "Outros" ao final da lista caso precise cadastrar alguém rápido
                    lista_fornecedores = nomes_banco + ["Outros"]
        except Exception:
            pass # Se der erro, mantêm a lista padrão ["Outros"]

    # Criamos abas para separar o cadastro Manual da Importacao
    aba1, aba2 = st.tabs(["Cadastro Manual", "Importacao em Massa (CSV)"])

    # --- ABA 1: CADASTRO MANUAL ---
    with aba1:
        st.info("O Valor Líquido será calculado: Custo + Impostos + Lucro")

        with st.form("form_cadastro"):
            st.subheader("1. Identificacao")
            col_id_1, col_id_2 = st.columns(2)
            
            with col_id_1:
                id_sku = st.text_input("SKU / Código Interno (Obrigatório)")
                descricao = st.text_input("Descricao do Produto (Obrigatório)")
                marca = st.text_input("Marca / Fabricante")
                
            with col_id_2:
                categoria = st.selectbox("Categoria", ["Geral", "Eletrônicos", "Vestuário", "Ferramentas", "Outros"])
                
                # --- AQUI ESTÁ A MUDANÇA: O SELECTBOX USA A LISTA DINÂMICA ---
                fornecedor = st.selectbox("Fornecedor", lista_fornecedores)
                
                c_est1, c_est2 = st.columns(2)
                estoque_atual = c_est1.number_input("Estoque Atual", min_value=0, step=1)
                estoque_minimo = c_est2.number_input("Estoque Mínimo", min_value=1, value=5)

            st.divider() 

            st.subheader("2. Custos e Precificacao")
            col_fin_1, col_fin_2, col_fin_3 = st.columns(3)
            with col_fin_1:
                preco_custo = st.number_input("Preco de Custo (R$)", min_value=0.00, step=0.01)
                lucro = st.number_input("Margem de Lucro (R$)", min_value=0.00, step=0.01)
            
            with col_fin_2:
                icms = st.number_input("ICMS (R$)", min_value=0.0, step=0.01)
                ipi = st.number_input("IPI (R$)", min_value=0.0, step=0.01)
            
            with col_fin_3:
                valor_st = st.number_input("ST (R$)", min_value=0.0, step=0.01)
                ncm = st.text_input("NCM")

            botao_salvar = st.form_submit_button("Salvar Produto")

        # --- LÓGICA DE SALVAR (MANUAL) ---
        if botao_salvar:
            # 1. Validacao de Campos Vazios
            erros = []
            if not id_sku: erros.append("O SKU é obrigatório.")
            if not descricao: erros.append("A Descricao é obrigatória.")
            
            # 2. TRAVA DE SKU DUPLICADO
            # Garante que dados_produtos esteja carregado (lógica global)
            try:
                lista_skus = dados_produtos["id_sku"].astype(str).tolist()
            except:
                lista_skus = [] # Se for o primeiro produto

            if str(id_sku) in lista_skus:
                erros.append(f"ERRO CRÍTICO: O SKU '{id_sku}' já existe no sistema!")

            if len(erros) > 0:
                for erro in erros:
                    st.error(erro)
            else:
                # Cálculo Automático
                valor_liquido = preco_custo + icms + ipi + valor_st + lucro
                
                # Criacao da Linha
                nova_linha = pd.DataFrame({
                    "id_sku": [id_sku],
                    "descricao": [descricao],
                    "categoria": [categoria],
                    "marca": [marca],
                    "fornecedor": [fornecedor],
                    "ncm": [ncm],
                    "preco_custo": [preco_custo],
                    "lucro": [lucro],
                    "icms": [icms], "ipi": [ipi], "st": [valor_st],
                    "valor_liquido": [valor_liquido],
                    "estoque_atual": [estoque_atual],
                    "estoque_minimo": [estoque_minimo],
                    "ativo": [True],
                    "data_cadastro": [datetime.now().strftime("%d/%m/%Y")]
                })

                # Salvar
                dados_produtos = pd.concat([dados_produtos, nova_linha], ignore_index=True)
                # Adicionei encoding='latin1' aqui para evitar o erro de Unicode no futuro
                dados_produtos.to_csv("Base de Dados.csv", sep=";", index=False, encoding="latin1")
                st.success(f"Produto {id_sku} cadastrado com sucesso!")
                st.rerun() # Atualiza a tela para limpar

    # --- ABA 2: IMPORTAcaO EM MASSA (CSV) ---
    with aba2:
        st.header("Importar Produtos via CSV")
        
        # --- SEcaO DE INSTRUcÕES ---
        with st.expander("📖 Leia as instrucões antes de importar", expanded=False):
            st.markdown("""
            Para que a importacao funcione corretamente, seu arquivo deve seguir estas regras:
            1. **Separador:** O arquivo deve ser salvo no formato **CSV (Separado por ponto e vírgula `;`)**.
            2. **Cabecalhos:** A primeira linha deve conter exatamente os nomes das colunas (id_sku, descricao, etc).
            3. **Decimais:** Em valores de dinheiro, utilize o **ponto (.)** como separador decimal.
            4. **SKUs Únicos:** Se um SKU do arquivo já existir no banco de dados, essa linha será ignorada.
            """)
            
            st.write("### Exemplo de preenchimento:")
            exemplo_dados = {
                "id_sku": ["CEL-S23-01", "CAM-IPH-15"],
                "descricao": ["Smartphone S23", "iPhone 15 Pro"],
                "fornecedor": ["Samsung", "Apple"],
                "preco_custo": [4500.00, 7200.00],
                "lucro": [500.00, 800.00],
                "valor_liquido": [5000.00, 8000.00]
            }
            st.table(exemplo_dados)
        
        st.divider()
        
        # --- ÁREA DE UPLOAD ---
        arquivo_upload = st.file_uploader("Arraste seu arquivo CSV aqui", type=["csv"])
        
        if arquivo_upload is not None:
            try:
                # Lendo o arquivo subido
                try:
                    df_novo = pd.read_csv(arquivo_upload, sep=";")
                except UnicodeDecodeError:
                    df_novo = pd.read_csv(arquivo_upload, sep=";", encoding="latin1")
                
                st.write("🔍 **Pré-visualizacao dos dados detectados:**")
                st.dataframe(df_novo.head()) 
                
                if st.button("Confirmar Importacao"):
                    # Verificacao de SKUs
                    skus_existentes = dados_produtos["id_sku"].astype(str).tolist()
                    
                    # Filtra apenas o que é novo
                    df_novo_filtrado = df_novo[~df_novo["id_sku"].astype(str).isin(skus_existentes)]
                    
                    qtd_total = len(df_novo)
                    qtd_novos = len(df_novo_filtrado)
                    qtd_ignorados = qtd_total - qtd_novos
                    
                    if qtd_novos > 0:
                        if "data_cadastro" not in df_novo_filtrado.columns:
                            df_novo_filtrado["data_cadastro"] = datetime.now().strftime("%d/%m/%Y")
                        
                        # Concatena e salva
                        dados_produtos = pd.concat([dados_produtos, df_novo_filtrado], ignore_index=True)
                        # Salva com encoding seguro
                        dados_produtos.to_csv("Base de Dados.csv", sep=";", index=False, encoding="latin1")
                        
                        st.success(f"✅ Sucesso! {qtd_novos} novos produtos adicionados.")
                        if qtd_ignorados > 0:
                            st.warning(f"⚠️ {qtd_ignorados} produtos ignorados (SKU repetido).")
                        
                        # Recarrega a página após importar
                        st.rerun()
                    else:
                        st.error("❌ Todos os produtos deste arquivo já existem no banco de dados.")
                        
            except Exception as e:
                st.error(f"Erro ao processar o arquivo: {e}")
#Tela de Consulta de Produtos
elif pagina == "Consultar Produto":
    st.title("Consulta de Produtos")

    # Criamos duas colunas: uma estreita para os filtros e uma larga para o resultado
    col_filtros, col_resultado = st.columns([1, 3])

    with col_filtros:
        st.subheader("Filtros de Busca")
        # Busca por SKU (Texto exato ou parcial)
        filtro_sku = st.text_input("Código SKU")
        
        # Busca por Descricao (Palavra-chave)
        filtro_desc = st.text_input("Descricao do Produto")
        
        st.divider()
        st.caption("Dica: A busca por descricao encontra palavras parciais (ex: 'azul' encontra 'Camiseta Azul').")

    with col_resultado:
        # Criamos uma cópia dos dados para nao alterar o original durante o filtro
        df_filtrado = dados_produtos.copy()

        # Lógica de Filtro em Tempo Real
        if filtro_sku:
            # Filtra por SKU (transformando tudo em string para evitar erro)
            df_filtrado = df_filtrado[df_filtrado["id_sku"].astype(str).str.contains(filtro_sku, case=False, na=False)]
        
        if filtro_desc:
            # Filtra por Descricao (case=False ignora maiúsculas/minúsculas)
            df_filtrado = df_filtrado[df_filtrado["descricao"].astype(str).str.contains(filtro_desc, case=False, na=False)]

        # Selecionamos apenas as colunas solicitadas
        # Nota: Use os nomes exatos das colunas do seu CSV aqui
        colunas_exibicao = [
            "id_sku", 
            "descricao", 
            "fornecedor", 
            "preco_custo", 
            "lucro", 
            "valor_liquido"
        ]
        
        # Verificar se as colunas existem antes de exibir (para evitar erro de arquivo vazio)
        try:
            exibicao = df_filtrado[colunas_exibicao]
            
            # Renomear apenas para ficar bonito na tabela do usuário
            exibicao.columns = ["SKU", "DESCRIcaO", "FORNECEDOR", "CUSTO (R$)", "LUCRO (R$)", "VALOR LÍQUIDO (R$)"]
            
            st.subheader(f"Resultados ({len(exibicao)} encontrados)")
            
            if len(exibicao) > 0:
                st.dataframe(
                    exibicao, 
                    use_container_width=True, 
                    hide_index=True # Esconde aquela coluna de números 0, 1, 2...
                )
            else:
                st.warning("Nenhum produto encontrado com esses termos.")
                
        except KeyError as e:
            st.error(f"Erro: Alguma coluna nao foi encontrada no CSV: {e}")
# --- 5. TELA DE CADASTRO DE PESSOAS ---
elif pagina == "Cadastrar Pessoa":
    st.title("Cadastro de Clientes e Fornecedores")

    # Tenta carregar a base de pessoas, se nao existir, cria uma vazia
    try:
        dados_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
    except:
        # Se o arquivo nao existir, cria o DataFrame com o cabecalho que definimos
        dados_pessoas = pd.DataFrame(columns=[
            "id_documento", "tipo_pessoa", "nome_razao", "nome_fantasia", 
            "rg_ie", "email", "telefone", "cep", "endereco", "numero", 
            "complemento", "bairro", "cidade", "uf", "categoria", 
            "limite_credito", "status", "data_cadastro"
        ])

    with st.form("form_pessoas"):
        st.subheader("1. Identificacao Principal")
        col_id_1, col_id_2, col_id_3 = st.columns([2, 2, 2])
        
        with col_id_1:
            tipo_pessoa = st.selectbox("Tipo de Pessoa", ["Física", "Jurídica"])
            # Muda o rótulo do campo conforme a escolha
            label_doc = "CPF (Obrigatório)" if tipo_pessoa == "Física" else "CNPJ (Obrigatório)"
            id_documento = st.text_input(label_doc)
            
        with col_id_2:
            categoria = st.selectbox("Categoria", ["Cliente", "Fornecedor", "Transportadora", "Ambos"])
            status = st.selectbox("Status Inicial", ["Ativo", "Inativo", "Bloqueado"])
            
        with col_id_3:
            limite_credito = st.number_input("Limite de Crédito (R$)", min_value=0.0, step=100.0)

        st.divider()

        st.subheader("2. Dados Pessoais / Empresariais")
        col_dados_1, col_dados_2 = st.columns(2)
        
        with col_dados_1:
            label_nome = "Nome Completo" if tipo_pessoa == "Física" else "Razao Social"
            nome_razao = st.text_input(label_nome)
            nome_fantasia = st.text_input("Nome Fantasia (Se houver)")
            
        with col_dados_2:
            label_rg = "RG" if tipo_pessoa == "Física" else "Inscricao Estadual"
            rg_ie = st.text_input(label_rg)
            email = st.text_input("E-mail para contato/NFe")
            telefone = st.text_input("WhatsApp / Telefone")

        st.divider()

        st.subheader("3. Endereco")
        col_end_1, col_end_2, col_end_3 = st.columns([1, 2, 1])
        with col_end_1:
            cep = st.text_input("CEP")
            numero = st.text_input("Número")
        with col_end_2:
            endereco = st.text_input("Logradouro (Rua/Av)")
            complemento = st.text_input("Complemento")
        with col_end_3:
            bairro = st.text_input("Bairro")
            cidade = st.text_input("Cidade")
            uf = st.selectbox("UF", ["AC", "AL", "AP", "AM", "BA", "CE", "DF", "ES", "GO", "MA", "MT", "MS", "MG", "PA", "PB", "PR", "PE", "PI", "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO"])

        botao_salvar_pessoa = st.form_submit_button("Finalizar Cadastro")

    # --- LÓGICA DE SALVAR PESSOA ---
    if botao_salvar_pessoa:
        erros_pessoa = []
        
        # 1. Validacao de Documento e Nome
        if not id_documento: erros_pessoa.append("O campo CPF/CNPJ é obrigatório.")
        if not nome_razao: erros_pessoa.append(f"O campo {label_nome} é obrigatório.")
        
        # 2. Trava de Duplicidade
        if str(id_documento) in dados_pessoas["id_documento"].astype(str).tolist():
            erros_pessoa.append(f"Este documento ({id_documento}) já está cadastrado no sistema!")

        if len(erros_pessoa) > 0:
            for erro in erros_pessoa:
                st.error(erro)
        else:
            # 3. Criar nova linha
            nova_pessoa = pd.DataFrame({
                "id_documento": [id_documento],
                "tipo_pessoa": [tipo_pessoa],
                "nome_razao": [nome_razao],
                "nome_fantasia": [nome_fantasia],
                "rg_ie": [rg_ie],
                "email": [email],
                "telefone": [telefone],
                "cep": [cep],
                "endereco": [endereco],
                "numero": [numero],
                "complemento": [complemento],
                "bairro": [bairro],
                "cidade": [cidade],
                "uf": [uf],
                "categoria": [categoria],
                "limite_credito": [limite_credito],
                "status": [status],
                "data_cadastro": [datetime.now().strftime("%d/%m/%Y")]
            })

            # 4. Salvar no CSV
            dados_pessoas = pd.concat([dados_pessoas, nova_pessoa], ignore_index=True)
            dados_pessoas.to_csv("Base_Pessoas.csv", sep=";", index=False)
            
            st.success(f"✅ {tipo_pessoa} '{nome_razao}' cadastrada com sucesso!")
# --- 6. TELA DE CONSULTA DE PESSOAS ---
elif pagina == "Consultar Pessoa":
    st.title("Consulta de Clientes / Fornecedores")

    # 1. Carregar os dados
    try:
        dados_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
    except:
        st.warning("Nenhuma base de pessoas encontrada. Cadastre alguém primeiro!")
        st.stop() # Para a execucao aqui caso nao haja arquivo

    # 2. Layout de Colunas
    col_filtros, col_resultado = st.columns([1, 3])

    with col_filtros:
        st.subheader("Filtros")
        filtro_doc = st.text_input("Buscar por CPF/CNPJ")
        filtro_nome = st.text_input("Buscar por Nome/Razao")
        
        filtro_cat = st.multiselect(
            "Filtrar Categoria", 
            ["Cliente", "Fornecedor", "Transportadora"],
            default=[]
        )
        
        st.divider()
        st.caption("A busca por nome funciona com termos parciais.")

    with col_resultado:
        # Criamos a cópia para filtrar
        df_p_filtrado = dados_pessoas.copy()

        # Lógica de Filtro em Tempo Real
        if filtro_doc:
            df_p_filtrado = df_p_filtrado[df_p_filtrado["id_documento"].astype(str).str.contains(filtro_doc, na=False)]
        
        if filtro_nome:
            df_p_filtrado = df_p_filtrado[df_p_filtrado["nome_razao"].astype(str).str.contains(filtro_nome, case=False, na=False)]
        
        if filtro_cat:
            # Filtra se a categoria está na lista selecionada no multiselect
            df_p_filtrado = df_p_filtrado[df_p_filtrado["categoria"].isin(filtro_cat)]

        # Selecao de Colunas para a Tabela (O que o usuário precisa ver rápido)
        colunas_ver = [
            "id_documento",
            "nome_razao",
            "categoria",
            "email",
            "telefone",
            "cidade",
            "status"
        ]

        try:
            exibicao_p = df_p_filtrado[colunas_ver]
            
            # Renomeando para ficar apresentável
            exibicao_p.columns = ["DOCUMENTO", "NOME / RAZaO SOCIAL", "CATEGORIA", "E-MAIL", "CONTATO", "CIDADE", "STATUS"]

            st.subheader(f"Registros Encontrados ({len(exibicao_p)})")
            
            if len(exibicao_p) > 0:
                st.dataframe(
                    exibicao_p, 
                    use_container_width=True, 
                    hide_index=True
                )
                
                # Widget extra: Ver detalhes completos
                if len(exibicao_p) == 1:
                    st.info("💡 Apenas um registro encontrado. Você pode ver todos os dados dele na tabela acima arrastando a barra de rolagem.")
            else:
                st.info("Nenhuma pessoa encontrada com esses critérios.")
                
        except KeyError as e:
            st.error(f"Erro nas colunas do arquivo: {e}")

# --- 7. TELA DE PEDIDOS (VERSAO ULTRA COMPATIVEL) ---
elif pagina == "Criar Pedido":
    st.title("Central de Pedidos")

    import unicodedata

    # Função para remover acentos, cedilhas e caracteres incompatíveis
    def limpar_texto(txt):
        if not txt or txt == 'nan': return ""
        # 1. Substituir Enters pelo seu separador escolhido
        txt = txt.replace("\n", "  |  ").replace("\r", "  |  ")
        # 2. Normalizar e remover acentos/cedilhas
        nfkd_form = unicodedata.normalize('NFKD', txt)
        txt = "".join([c for c in nfkd_form if not unicodedata.combining(c)])
        # 3. Remover o ponto e vírgula para não confundir com o separador do CSV
        txt = txt.replace(";", ",")
        return txt.strip()

    def ler_csv_safe(caminho):
        try:
            return pd.read_csv(caminho, sep=";")
        except UnicodeDecodeError:
            return pd.read_csv(caminho, sep=";", encoding="latin-1")

    # --- INICIALIZACAO DE ESTADOS ---
    if "carrinho" not in st.session_state:
        st.session_state.carrinho = []
    if "cliente_selecionado" not in st.session_state:
        st.session_state.cliente_selecionado = None
    if "produto_selecionado" not in st.session_state:
        st.session_state.produto_selecionado = None

    # Logica de ID Sequencial
    if os.path.exists("Base_Pedido.csv"):
        try:
            base_pedidos_temp = ler_csv_safe("Base_Pedido.csv")
            proximo_id = base_pedidos_temp["id_pedido"].max() + 1 if not base_pedidos_temp.empty else 1
        except:
            proximo_id = 1
    else:
        proximo_id = 1

    st.subheader(f"Pedido Nº: {proximo_id}")

    # --- FUNCOES DE BUSCA (DIALOGS) ---
    @st.dialog("Buscar Cliente")
    def buscar_cliente_pop():
        st.write("Pesquise e selecione o cliente.")
        filtro = st.text_input("Nome ou CPF/CNPJ")
        if os.path.exists("Base_Pessoas.csv"):
            df_p = ler_csv_safe("Base_Pessoas.csv")
            if filtro:
                df_p = df_p[df_p["nome_razao"].str.contains(filtro, case=False, na=False) | 
                            df_p["id_documento"].astype(str).str.contains(filtro, na=False)]
            
            for _, row in df_p.head(10).iterrows():
                col1, col2 = st.columns([3, 1])
                col1.write(f"**{row['nome_razao']}** ({row['id_documento']})")
                if col2.button("Selecionar", key=f"sel_p_{row['id_documento']}"):
                    st.session_state.cliente_selecionado = row.to_dict()
                    st.rerun()
        else:
            st.error("Base de Pessoas nao encontrada!")

    @st.dialog("Buscar Produto")
    def buscar_produto_pop():
        st.write("Pesquise o SKU ou Descricao")
        filtro = st.text_input("Palavra-chave")
        if os.path.exists("Base de Dados.csv"):
            df_prod = ler_csv_safe("Base de Dados.csv")
            if filtro:
                df_prod = df_prod[df_prod["descricao"].str.contains(filtro, case=False, na=False) | 
                                  df_prod["id_sku"].astype(str).str.contains(filtro, na=False)]
            
            for _, row in df_prod.head(10).iterrows():
                col1, col2 = st.columns([3, 1])
                col1.write(f"**{row['id_sku']}** - {row['descricao']}")
                if col2.button("Selecionar", key=f"sel_prod_{row['id_sku']}"):
                    st.session_state.produto_selecionado = row.to_dict()
                    st.rerun()
        else:
            st.error("Base de Produtos nao encontrada!")

    # --- ÁREA 1: IDENTIFICACAO DO CLIENTE ---
    with st.container(border=True):
        col_cli_1, col_cli_2 = st.columns([3, 1])
        with col_cli_1:
            doc_exibicao = st.session_state.cliente_selecionado['id_documento'] if st.session_state.cliente_selecionado else ""
            st.text_input("Cliente (Selecione na busca)", value=doc_exibicao, disabled=True)
        with col_cli_2:
            st.write("##")
            if st.button("Buscar Cliente", use_container_width=True):
                buscar_cliente_pop()

        if st.session_state.cliente_selecionado:
            c = st.session_state.cliente_selecionado
            st.success(f"✅ **{c['nome_razao']}** | {c['cidade']}-{c['uf']} | Limite: R$ {c['limite_credito']}")

    # --- ÁREA 2: INCLUSAO DE PRODUTOS ---
    with st.container(border=True):
        st.write("### Adicionar Itens")
        col_prod_1, col_prod_2, col_prod_3 = st.columns([2, 1, 1])
        
        with col_prod_1:
            sku_exibicao = st.session_state.produto_selecionado['id_sku'] if st.session_state.produto_selecionado else ""
            st.text_input("SKU (Selecione na busca)", value=sku_exibicao, disabled=True)
        with col_prod_2:
            st.write("##")
            if st.button("Buscar Produto", use_container_width=True):
                buscar_produto_pop()
        with col_prod_3:
            qtd = st.number_input("Quantidade", min_value=1, value=1)

        if st.session_state.produto_selecionado:
            p = st.session_state.produto_selecionado
            st.info(f"**Produto:** {p['descricao']} | **Preco Base:** R$ {p['valor_liquido']:.2f}")
            
            col_desc_1, col_custo, col_final = st.columns([2, 1, 1])
            
            desconto = col_desc_1.number_input(
                "Desconto (R$) - Use negativo para Acréscimo", 
                value=0.0, 
                step=1.0
            )
            
            valor_final_item = float(p['valor_liquido']) - desconto
            preco_custo = float(p.get('preco_custo', 0.0))
            
            col_custo.metric("Preco de Custo", f"R$ {preco_custo:.2f}")

            label_delta = "Acréscimo" if desconto < 0 else "Desconto"
            col_final.metric("Preco Unitário Final", f"R$ {valor_final_item:.2f}", delta=f"{-desconto:.2f} ({label_delta})")

            if st.button("Adicionar ao Carrinho ✅", use_container_width=True):
                if any(item['sku'] == p['id_sku'] for item in st.session_state.carrinho):
                    st.warning("Este produto já está no carrinho!")
                else:
                    st.session_state.carrinho.append({
                        "sku": p['id_sku'],
                        "descricao": p['descricao'],
                        "qtd": qtd,
                        "valor_unit": valor_final_item,
                        "subtotal": valor_final_item * qtd
                    })
                    st.session_state.produto_selecionado = None
                    st.rerun()

    # --- ÁREA 3: REVISAO E FINALIZACAO ---
    if st.session_state.carrinho:
        st.divider()
        st.subheader("🛒 Resumo do Pedido")
        df_carrinho = pd.DataFrame(st.session_state.carrinho)
        st.dataframe(df_carrinho, use_container_width=True, hide_index=True)
        
        if st.button("Esvaziar Carrinho"):
            st.session_state.carrinho = []
            st.rerun()

        subtotal_itens = df_carrinho["subtotal"].sum()

        with st.form("finalizar_venda"):
            f1, f2, f3 = st.columns([2, 1, 1])
            
            with f1:
                tipo = st.selectbox("Tipo", ["ORCAMENTO", "PEDIDO", "COTACAO"])
                obs = st.text_area("Observacoes")
            with f2:
                frete = st.number_input("Frete (R$)", min_value=0.0, step=5.0)
                st.write(f"**Subtotal:** R$ {subtotal_itens:.2f}")
            with f3:
                total_final = subtotal_itens + frete
                st.metric("TOTAL GERAL", f"R$ {total_final:.2f}")

            if st.form_submit_button("CONFIRMAR E SALVAR ✅", use_container_width=True):
                if not st.session_state.cliente_selecionado:
                    st.error("Selecione um cliente!")
                else:
                    # APLICANDO A LIMPEZA PESADA NAS OBSERVACOES E NOMES
                    obs_limpa = limpar_texto(obs)
                    cliente_limpo = limpar_texto(st.session_state.cliente_selecionado['nome_razao'])
                    
                    novas_linhas = []
                    for item in st.session_state.carrinho:
                        novas_linhas.append({
                            "id_pedido": proximo_id,
                            "data_pedido": datetime.now().strftime("%d/%m/%Y %H:%M"),
                            "doc_cliente": st.session_state.cliente_selecionado['id_documento'],
                            "nome_cliente": cliente_limpo,
                            "sku_item": item['sku'],
                            "qtd": item['qtd'],
                            "valor_final": item['valor_unit'],
                            "frete_total": frete,
                            "tipo": tipo,
                            "observacao": obs_limpa
                        })
                    
                    df_save = pd.DataFrame(novas_linhas)
                    header_status = not os.path.exists("Base_Pedido.csv")
                    df_save.to_csv("Base_Pedido.csv", mode='a', sep=";", index=False, header=header_status, encoding="latin-1")
                    
                    st.success("✅ Pedido Gravado com Sucesso!")
                    st.session_state.carrinho = []
                    st.session_state.cliente_selecionado = None
                    st.rerun()
# --- 8. TELA DE CONSULTA DE PEDIDOS (COM EDIcaO VIA POPUP) ---
## --- 8. TELA DE CONSULTA DE PEDIDOS (VERSÃO AJUSTADA E ROBUSTA) ---
elif pagina == "Consultar Pedido":
    st.title("🔎 Gestao e Consulta de Pedidos")

    import unicodedata

    # Função de limpeza repetida para garantir consistência no salvamento da edição
    def limpar_texto_local(txt):
        if not txt or str(txt).lower() == 'nan': return ""
        txt = str(txt).replace("\n", "  |  ").replace("\r", "  |  ")
        nfkd_form = unicodedata.normalize('NFKD', txt)
        txt = "".join([c for c in nfkd_form if not unicodedata.combining(c)])
        txt = txt.replace(";", ",")
        return txt.strip()

    # 1. Carregamento e Padronizacao com tratamento de erro numérico
    try:
        # Lendo com latin-1 para suportar caracteres do Windows/Excel
        df_pedidos = pd.read_csv("Base_Pedido.csv", sep=";", encoding="latin-1")
        df_produtos = pd.read_csv("Base de Dados.csv", sep=";", encoding="latin-1")
        df_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";", encoding="latin-1")
        
        # --- CORREÇÃO DE TIPOS (Evita o erro de soma str + float) ---
        df_pedidos['valor_final'] = pd.to_numeric(df_pedidos['valor_final'], errors='coerce').fillna(0.0)
        df_pedidos['qtd'] = pd.to_numeric(df_pedidos['qtd'], errors='coerce').fillna(0).astype(int)
        df_pedidos['frete_total'] = pd.to_numeric(df_pedidos['frete_total'], errors='coerce').fillna(0.0)
        
        df_pedidos['sku_item'] = df_pedidos['sku_item'].astype(str)
        df_produtos['id_sku'] = df_produtos['id_sku'].astype(str)
        df_pedidos['doc_cliente'] = df_pedidos['doc_cliente'].astype(str)
        
        # Tratamento de data
        df_pedidos['data_pedido_dt'] = pd.to_datetime(df_pedidos['data_pedido'], format="%d/%m/%Y %H:%M", errors='coerce')
        
    except Exception as e:
        st.error(f"Erro ao carregar bases: {e}")
        st.stop()

    # 2. Layout Lateral (Filtros)
    col_filtros, col_detalhe = st.columns([1, 2.5])

    with col_filtros:
        st.subheader("Filtros")
        f_id_check = st.checkbox("Número Pedido")
        f_id = st.number_input("ID", min_value=1, step=1, disabled=not f_id_check)
        
        df_f = df_pedidos.copy()
        if f_id_check: 
            df_f = df_f[df_f["id_pedido"] == f_id]

        st.divider()
        lista_ids = df_f["id_pedido"].unique()
        id_selecionado = st.selectbox("Selecione o Pedido", sorted(lista_ids, reverse=True)) if len(lista_ids) > 0 else None

    # 3. Coluna de Detalhes / Edicao
    with col_detalhe:
        if id_selecionado:
            if "modo_edicao" not in st.session_state:
                st.session_state.modo_edicao = False

            itens_venda = df_pedidos[df_pedidos["id_pedido"] == id_selecionado]
            
            # --- MODO EDIÇÃO ---
            if st.session_state.modo_edicao:
                st.warning("⚠️ MODO DE EDIÇÃO ATIVADO")
                
                if "edit_carrinho" not in st.session_state:
                    st.session_state.edit_carrinho = []
                    for _, r in itens_venda.iterrows():
                        st.session_state.edit_carrinho.append({
                            "sku": r["sku_item"], 
                            "qtd": int(r["qtd"]), 
                            "valor_unit": float(r["valor_final"])
                        })
                    st.session_state.edit_frete = float(itens_venda.iloc[0]["frete_total"])
                    st.session_state.edit_tipo = itens_venda.iloc[0]["tipo"]
                    # Volta o " | " para "Enter" na área de texto para facilitar a edição
                    obs_original = str(itens_venda.iloc[0]["observacao"]).replace("  |  ", "\n")
                    st.session_state.edit_obs = obs_original

                ce1, ce2, ce3 = st.columns([1, 1, 1])
                novo_tipo = ce1.selectbox("Tipo", ["ORCAMENTO", "PEDIDO", "COTACAO"], 
                                         index=["ORCAMENTO", "PEDIDO", "COTACAO"].index(st.session_state.edit_tipo) if st.session_state.edit_tipo in ["ORCAMENTO", "PEDIDO", "COTACAO"] else 0)
                novo_frete = ce2.number_input("Frete (R$)", value=st.session_state.edit_frete)
                
                if ce3.button("❌ CANCELAR", use_container_width=True):
                    st.session_state.modo_edicao = False
                    if "edit_carrinho" in st.session_state: del st.session_state.edit_carrinho
                    st.rerun()

                nova_obs = st.text_area("Observacao", value=st.session_state.edit_obs if st.session_state.edit_obs != 'nan' else "")

                st.write("### Itens do Pedido")
                for i, item in enumerate(st.session_state.edit_carrinho):
                    c_i1, c_i2, c_i3, c_i4 = st.columns([2, 1, 1, 0.5])
                    desc_p = df_produtos[df_produtos['id_sku'] == item['sku']]['descricao'].values[0] if item['sku'] in df_produtos['id_sku'].values else "Produto nao encontrado"
                    
                    c_i1.caption(f"{item['sku']} - {desc_p}")
                    item['qtd'] = c_i2.number_input(f"Qtd", value=int(item['qtd']), key=f"q_ed_{i}", min_value=1)
                    item['valor_unit'] = c_i3.number_input(f"R$", value=float(item['valor_unit']), key=f"v_ed_{i}")
                    if c_i4.button("❌", key=f"del_{i}"):
                        st.session_state.edit_carrinho.pop(i)
                        st.rerun()

                if st.button("SALVAR ALTERAÇÕES ✅", use_container_width=True, type="primary"):
                    # Remove o pedido antigo e insere os itens novos
                    df_base_limpa = df_pedidos[df_pedidos["id_pedido"] != id_selecionado].copy()
                    origem = itens_venda.iloc[0]
                    
                    obs_limpa = limpar_texto_local(nova_obs)
                    
                    novas_linhas = []
                    for it in st.session_state.edit_carrinho:
                        novas_linhas.append({
                            "id_pedido": id_selecionado, 
                            "data_pedido": origem["data_pedido"],
                            "doc_cliente": origem["doc_cliente"], 
                            "nome_cliente": origem["nome_cliente"],
                            "sku_item": it["sku"], 
                            "qtd": it["qtd"], 
                            "valor_final": it["valor_unit"],
                            "frete_total": novo_frete, 
                            "tipo": novo_tipo, 
                            "observacao": obs_limpa
                        })
                    
                    df_save = pd.concat([df_base_limpa, pd.DataFrame(novas_linhas)], ignore_index=True)
                    df_save.to_csv("Base_Pedido.csv", sep=";", index=False, encoding="latin-1")
                    
                    st.session_state.modo_edicao = False
                    if "edit_carrinho" in st.session_state: del st.session_state.edit_carrinho
                    st.success("Pedido atualizado com sucesso!")
                    st.rerun()

            # --- MODO LEITURA (CONSULTA) ---
            else:
                itens_completos = itens_venda.merge(df_produtos[['id_sku', 'descricao', 'marca']], left_on='sku_item', right_on='id_sku', how='left')
                
                # Cálculos protegidos contra erros de tipo
                frete = float(itens_venda.iloc[0]["frete_total"])
                soma_itens = (itens_completos["valor_final"] * itens_completos["qtd"]).sum()
                total = soma_itens + frete

                with st.container(border=True):
                    c_h1, c_h2 = st.columns([2, 1])
                    c_h1.markdown(f"### {itens_venda.iloc[0]['nome_cliente']}")
                    if c_h2.button("📝 EDITAR PEDIDO", use_container_width=True):
                        st.session_state.modo_edicao = True
                        st.rerun()
                    
                    st.divider()
                    m1, m2, m3 = st.columns(3)
                    m1.metric("Frete", f"R$ {frete:.2f}")
                    m2.metric("Subtotal Itens", f"R$ {soma_itens:.2f}")
                    m3.subheader(f"Total: R$ {total:.2f}")

                st.write("**Itens do Pedido:**")
                st.dataframe(itens_completos[['sku_item', 'descricao', 'qtd', 'valor_final']], 
                             use_container_width=True, hide_index=True)
                
                obs = itens_venda.iloc[0]['observacao']
                if pd.notna(obs) and str(obs).lower() != 'nan' and str(obs).strip() != "":
                    # Mostra a observação com "Enter" na tela de consulta para ficar bonito
                    st.info(f"**Observações:**\n\n{str(obs).replace('  |  ', '\n')}")
        else:
            st.info("💡 Selecione um pedido na lista à esquerda para ver os detalhes.")
### --- 9. TELA DE FORMALIZAcaO DE PROPOSTA (CORRIGIDA) --- 
elif pagina == "Formalizacao":
    st.title("Formalizacao de Proposta")
    
    # 1. Carregamento e Padronizacao
    try:
        df_pedidos = pd.read_csv("Base_Pedido.csv", sep=";")
        df_produtos = pd.read_csv("Base de Dados.csv", sep=";")
        df_pessoas = pd.read_csv("Base_Pessoas.csv", sep=";")
        
        # Conversao de tipos para garantir o cruzamento (merge)
        df_pedidos['id_pedido'] = df_pedidos['id_pedido'].astype(int)
        df_pedidos['sku_item'] = df_pedidos['sku_item'].astype(str)
        df_produtos['id_sku'] = df_produtos['id_sku'].astype(str)
        df_pessoas['id_documento'] = df_pessoas['id_documento'].astype(str)
        
        lista_pedidos = sorted(df_pedidos["id_pedido"].unique(), reverse=True)
    except Exception as e:
        st.error(f"Erro ao carregar bases: {e}")
        st.stop()

    # 2. Selecao do Pedido
    id_escolhido = st.selectbox("Selecione o Número do Pedido", lista_pedidos, index=None, placeholder="Escolha um pedido...")

    if id_escolhido:
        # Filtra dados do Pedido
        dados_venda = df_pedidos[df_pedidos["id_pedido"] == id_escolhido]
        doc_cliente = str(dados_venda.iloc[0]["doc_cliente"])
        
        # AJUSTE NAS COLUNAS: Usando apenas o que existe no seu CSV (id_sku, descricao, marca, preco_custo)
        itens_completos = dados_venda.merge(
            df_produtos[['id_sku', 'descricao', 'marca', 'preco_custo']], 
            left_on='sku_item', right_on='id_sku', how='left'
        )
        
        # Dados do Cliente
        cliente_info = df_pessoas[df_pessoas["id_documento"] == doc_cliente].iloc[0]

        # --- PAINEL DE CONFERÊNCIA ---
        with st.container(border=True):
            st.subheader(f"Resumo: Pedido #{id_escolhido}")
            c1, c2 = st.columns(2)
            with c1:
                st.write(f"**Razao Social:** {cliente_info['nome_razao']}")
                st.write(f"**Data:** {dados_venda.iloc[0]['data_pedido']}")
                st.caption(f"📍 {cliente_info['endereco']}, {cliente_info['numero']} - {cliente_info['cidade']}/{cliente_info['uf']}")
            with c2:
                custo_total = (itens_completos['preco_custo'] * itens_completos['qtd']).sum()
                # Soma subtotal dos itens + o frete único do pedido
                venda_total = (itens_completos['valor_final'] * itens_completos['qtd']).sum() + float(dados_venda.iloc[0]['frete_total'])
                st.metric("Total Venda (c/ Frete)", f"R$ {venda_total:.2f}")
                st.write(f"**Custo Total Est.:** R$ {custo_total:.2f}")

            # Exibicao da Observacao do Pedido
            obs_pedido = dados_venda.iloc[0]['observacao']
            if pd.notna(obs_pedido) and str(obs_pedido).lower() != 'nan':
                st.warning(f"📝 **Observacao do Pedido:** {obs_pedido}")

            st.write("**Itens Selecionados:**")
            st.dataframe(itens_completos[['sku_item', 'descricao', 'marca', 'qtd', 'valor_final']], use_container_width=True, hide_index=True)

        st.divider()

        # 3. Inputs Manuais para o Documento
        st.subheader("Dados Adicionais para Proposta_Modelo")
        with st.form("form_formalizacao"):
            f1, f2 = st.columns(2)
            with f1:
                n_pregao = st.text_input("Nº do Pregao / Processo")
                validade = st.text_input("Validade da Proposta (ex: 60 dias)")
            with f2:
                prazo = st.text_input("Prazo de Entrega (ex: 15 dias úteis)")
                contato_doc = st.text_input("Pessoa de Contato", value=cliente_info['nome_razao'])
            
            especificacoes = st.text_area("Especificacões Técnicas Solicitadas")
            
            botao_gerar = st.form_submit_button("Gerar Proposta (Word)", use_container_width=True)

        # 4. Geracao do Word
        if botao_gerar:
            if not all([n_pregao, validade, prazo, especificacoes]):
                st.error("Preencha todos os campos obrigatórios para gerar o documento.")
            else:
                try:
                    from num2words import num2words
                    doc = Document("Proposta_Modelo.docx")
                    
                    # Valor por extenso
                    valor_extenso = num2words(venda_total, lang='pt_BR', to='currency').upper()

                    # Substituicao de Tags
                    subs = {
                        "[Razao_UASG]": cliente_info['nome_razao'],
                        "[N_pregao]": n_pregao,
                        "[Esp_solicitadas]": especificacoes,
                        "[Validade_Proposta]": validade,
                        "[Prazo_entrega]": prazo,
                        "[Endereco_Cliente]": f"{cliente_info['endereco']}, {cliente_info['numero']}",
                        "[Contato_Cliente]": contato_doc,
                        "MIL QUINHENTOS E QUARENTA REAIS": valor_extenso
                    }

                    for p in doc.paragraphs:
                        for tag, val in subs.items():
                            if tag in p.text:
                                p.text = p.text.replace(tag, str(val))

                    # Preenchimento da Tabela
                    if doc.tables:
                        tabela = doc.tables[0]
                        for i, it in itens_completos.iterrows():
                            cells = tabela.add_row().cells
                            cells[0].text = str(i + 1)
                            cells[1].text = str(it['descricao'])
                            cells[2].text = str(it['marca'])
                            cells[3].text = "---" # PartNumber nao existe no seu CSV, deixamos fixo ou vazio
                            cells[4].text = str(it['qtd'])
                            cells[5].text = f"R$ {it['valor_final']:.2f}"
                            cells[6].text = f"R$ {(it['valor_final'] * it['qtd']):.2f}"

                    nome_final = f"Proposta_{id_escolhido}.docx"
                    doc.save(nome_final)
                    
                    with open(nome_final, "rb") as f:
                        st.download_button("Baixar Proposta Gerada", f, file_name=nome_final)
                    st.success("Documento gerado!")

                except Exception as e:
                    st.error(f"Erro na geracao do documento: {e}")

# --- TELA 0. DASHBOARD (VERSÃO INTEGRAL OTIMIZADA COM INDICADORES FINANCEIROS REAIS) ---
elif pagina == "Dashboard":
    st.title("BI e Dashboard de Vendas")

    def ler_csv_safe(c):
        try: return pd.read_csv(c, sep=";", encoding="latin-1")
        except: return pd.read_csv(c, sep=";")

    @st.dialog("Pesquisar Produto")
    def pesquisar_produto_dialog():
        termo = st.text_input("Digite o nome ou SKU:")
        if termo:
            res = df_produtos[df_produtos['descricao'].str.contains(termo, case=False, na=False)]
            if not res.empty:
                for _, r in res.head(10).iterrows():
                    if st.button(f"{r['id_sku']} - {r['descricao'][:50]}", key=f"db_p_{r['id_sku']}"):
                        st.session_state.sku_filtro_db = str(r['id_sku']); st.rerun()
            else: st.warning("Nenhum produto encontrado.")

    @st.dialog("Pesquisar Cliente")
    def pesquisar_pessoa_dialog():
        termo = st.text_input("Digite o nome ou Razão Social:")
        if termo:
            res = df_pess[df_pess['nome_razao'].str.contains(termo, case=False, na=False)]
            if not res.empty:
                for _, r in res.head(10).iterrows():
                    if st.button(f"{r['id_documento']} - {r['nome_razao']}", key=f"db_c_{r['id_documento']}"):
                        st.session_state.cnpj_filtro_db = str(r['id_documento']); st.rerun()
            else: st.warning("Nenhum cliente encontrado.")

    if not os.path.exists("Base_Pedido.csv"): st.info("Ainda não existem pedidos cadastrados.")
    else:
        df_pedidos = ler_csv_safe("Base_Pedido.csv")
        df_produtos = ler_csv_safe("Base de Dados.csv") if os.path.exists("Base de Dados.csv") else pd.DataFrame()
        df_pedidos['tipo'] = df_pedidos['tipo'].astype(str).str.upper().str.strip()
        df_pedidos['sku_item'] = df_pedidos['sku_item'].astype(str)
        df_pedidos['doc_cliente'] = df_pedidos['doc_cliente'].astype(str)

        if os.path.exists("Base_Pessoas.csv"):
            df_pess = ler_csv_safe("Base_Pessoas.csv")
            df_pess['id_documento'] = df_pess['id_documento'].astype(str)
            df_mestre = pd.merge(df_pedidos, df_pess[['id_documento', 'nome_razao', 'cidade', 'uf', 'tipo_pessoa']], left_on='doc_cliente', right_on='id_documento', how='left')
        else: df_mestre = df_pedidos.copy()

        if not df_produtos.empty:
            df_produtos['id_sku'] = df_produtos['id_sku'].astype(str)
            # Traz todas as colunas de produtos para garantir que 'preco_custo' e % de lucro venham
            cols_prod = [c for c in df_produtos.columns if c not in df_mestre.columns or c == 'id_sku']
            df_mestre = pd.merge(df_mestre, df_produtos[cols_prod], left_on='sku_item', right_on='id_sku', how='left')
        
        df_mestre['valor_final'] = pd.to_numeric(df_mestre['valor_final'], errors='coerce').fillna(0)
        df_mestre['qtd'] = pd.to_numeric(df_mestre['qtd'], errors='coerce').fillna(0)
        df_mestre['data_pedido_dt'] = pd.to_datetime(df_mestre['data_pedido'], dayfirst=True, errors='coerce')
        
        # --- CORREÇÃO: Faturamento do Item = Preço Unitário * Quantidade ---
        df_mestre['faturamento_item'] = df_mestre['valor_final'] * df_mestre['qtd']

        with st.expander("Filtros de Pesquisa Avançada", expanded=True):
            f_col1, f_col2, f_col3, f_col4 = st.columns(4, vertical_alignment="bottom")
            with f_col1:
                min_d = df_mestre['data_pedido_dt'].min() if not df_mestre['data_pedido_dt'].isnull().all() else datetime.now()
                max_d = df_mestre['data_pedido_dt'].max() if not df_mestre['data_pedido_dt'].isnull().all() else datetime.now()
                data_inicio = st.date_input("Data Início", min_d)
                data_fim = st.date_input("Data Fim", max_d)
            with f_col2:
                razao_social = st.text_input("Razão Social")
                c1, c2 = st.columns([4, 1], vertical_alignment="bottom")
                with c1: cnpj_f = st.text_input("CNPJ ou CPF", value=st.session_state.get('cnpj_filtro_db', ""))
                with c2: 
                    if st.button("🔍", key="btn_pess_db"): pesquisar_pessoa_dialog()
            with f_col3:
                tipo_p_f = st.selectbox("Tipo de Pessoa", ["Todos"] + list(df_mestre['tipo_pessoa'].dropna().unique()) if 'tipo_pessoa' in df_mestre else ["Todos"])
                s1, s2 = st.columns([4, 1], vertical_alignment="bottom")
                with s1: sku_f = st.text_input("SKU do Item", value=st.session_state.get('sku_filtro_db', ""))
                with s2:
                    if st.button("🔍", key="btn_sku_db"): pesquisar_produto_dialog()
                # ADICIONADO AQUI: Filtro de Categoria/Tipo de Pedido
                tipo_pedido_f = st.multiselect("Categoria / Tipo", list(df_mestre['tipo'].dropna().unique()))
            with f_col4:
                cidade_f = st.text_input("Cidade")
                estado_f = st.text_input("Estado (UF)")
                # Slider ajustado para o Faturamento real
                v_min, v_max = float(df_mestre['faturamento_item'].min() or 0), float(df_mestre['faturamento_item'].max() or 100)
                faixa_preco = st.slider("Faixa de Valor (R$)", v_min, v_max, (v_min, v_max))

        df_f = df_mestre[(df_mestre['data_pedido_dt'].dt.date >= data_inicio) & (df_mestre['data_pedido_dt'].dt.date <= data_fim) & (df_mestre['faturamento_item'] >= faixa_preco[0]) & (df_mestre['faturamento_item'] <= faixa_preco[1])]
        if razao_social: df_f = df_f[df_f['nome_razao'].str.contains(razao_social, case=False, na=False)]
        if cnpj_f: df_f = df_f[df_f['doc_cliente'].astype(str).str.contains(cnpj_f)]
        if cidade_f: df_f = df_f[df_f['cidade'].str.contains(cidade_f, case=False, na=False)]
        if estado_f: df_f = df_f[df_f['uf'].str.contains(estado_f, case=False, na=False)]
        if sku_f: df_f = df_f[df_f['sku_item'].str.contains(sku_f, case=False, na=False)]
        if tipo_p_f != "Todos": df_f = df_f[df_f['tipo_pessoa'] == tipo_p_f]
        if tipo_pedido_f: df_f = df_f[df_f['tipo'].isin(tipo_pedido_f)] # ADICIONADA A LÓGICA DO FILTRO AQUI

        # --- INDICADORES DE PERFORMANCE ---
        st.subheader("Indicadores de Performance")
        i1, i2, i3 = st.columns(3)
        
        fat_total = df_f['faturamento_item'].sum()
        qtd_total = df_f['qtd'].sum()
        ped_unicos = df_f['id_pedido'].nunique()
        
        # --- CORREÇÃO: Lógica de Custo e Lucro ---
        col_lucro = next((c for c in df_f.columns if 'lucro' in c.lower() or 'margem' in c.lower()), None)
        if col_lucro:
            # Se achar a coluna de %, aplica sobre o faturamento do item
            df_f[col_lucro] = pd.to_numeric(df_f[col_lucro], errors='coerce').fillna(0)
            mult = df_f[col_lucro].apply(lambda x: x / 100 if x > 1 else x)
            lucro = (df_f['faturamento_item'] * mult).sum()
            custo_total = fat_total - lucro
        elif 'preco_custo' in df_f:
            # Se não achar a %, usa o custo unitário multiplicado pela quantidade
            df_f['preco_custo'] = pd.to_numeric(df_f['preco_custo'], errors='coerce').fillna(0)
            custo_total = (df_f['preco_custo'] * df_f['qtd']).sum()
            lucro = fat_total - custo_total
        else:
            custo_total, lucro = 0, 0
        
        with i1:
            st.metric("Faturamento Total", f"R$ {fat_total:,.2f}")
            st.metric("Lucro Bruto", f"R$ {lucro:,.2f}")
            st.metric("Ticket Médio", f"R$ {(fat_total/ped_unicos if ped_unicos > 0 else 0):,.2f}")
        with i2:
            st.metric("Qtd Total Itens", f"{int(qtd_total)}")
            st.metric("Margem Bruta (%)", f"{(lucro/fat_total*100 if fat_total > 0 else 0):,.1f}%")
            st.metric("Clientes Únicos", f"{df_f['doc_cliente'].nunique()}")
        with i3:
            st.metric("Total de Pedidos", f"{ped_unicos}")
            st.metric("Custo Total", f"R$ {custo_total:,.2f}")
            st.metric("Média Itens/Pedido", f"{(qtd_total/ped_unicos if ped_unicos > 0 else 0):,.1f}")

        # --- GRÁFICOS ESTRATÉGICOS (AJUSTADOS PARA VERMELHO E SEM MARGENS) ---
        st.subheader("Análises Estratégicas")
        g1, g2 = st.columns(2)
        with g1:
            st.write("**1. Mix por Categoria (R$)**")
            fig_pie = px.pie(df_f, values='faturamento_item', names='tipo', hole=0.4, color_discrete_sequence=px.colors.sequential.Reds_r)
            fig_pie.update_layout(margin=dict(l=0, r=0, t=0, b=0))
            st.plotly_chart(fig_pie, use_container_width=True)
            
            st.write("**3. Top 5 Produtos (Qtd)**")
            if 'descricao' in df_f:
                top = df_f.groupby('descricao')['qtd'].sum().nlargest(5).reset_index()
                fig_bar1 = px.bar(top, x='qtd', y='descricao', orientation='h', color='qtd', color_continuous_scale='Reds')
                fig_bar1.update_layout(yaxis={'categoryorder':'total ascending'}, margin=dict(l=0, r=0, t=0, b=0))
                st.plotly_chart(fig_bar1, use_container_width=True)
                
        with g2:
            st.write("**2. Evolução Faturamento Diário**")
            ev = df_f.groupby(df_f['data_pedido_dt'].dt.date)['faturamento_item'].sum().reset_index()
            # Usando o vermelho padrão do Plotly (Coral) para a área
            fig_area = px.area(ev, x='data_pedido_dt', y='faturamento_item', markers=True, color_discrete_sequence=['#ef553b'])
            fig_area.update_layout(margin=dict(l=0, r=0, t=0, b=0))
            st.plotly_chart(fig_area, use_container_width=True)
            
            st.write("**4. Faturamento por Estado**")
            if 'uf' in df_f:
                uf = df_f.groupby('uf')['faturamento_item'].sum().reset_index()
                fig_bar2 = px.bar(uf, x='uf', y='faturamento_item', color='faturamento_item', color_continuous_scale='Reds')
                fig_bar2.update_layout(margin=dict(l=0, r=0, t=0, b=0))
                st.plotly_chart(fig_bar2, use_container_width=True)

        st.divider(); st.write("**Detalhamento**")
        st.dataframe(df_f.drop(columns=['data_pedido_dt', 'id_sku', 'id_documento'], errors='ignore'), use_container_width=True, hide_index=True)
        st.download_button("📥 Exportar Relatório", df_f.to_csv(index=False, sep=';', encoding='latin-1'), "bi.csv", "text/csv")